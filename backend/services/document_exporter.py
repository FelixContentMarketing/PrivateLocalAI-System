import io
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_txt(text: str, document_type: str) -> bytes:
    """Export generated text as UTF-8 plain text."""
    return text.encode("utf-8")


def export_docx(text: str, document_type: str, metadata: dict | None = None) -> bytes:
    """Export generated text as DOCX with professional formatting."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run("ENTWURF")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = "Calibri"

    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sub_header.add_run("Erstellt mit PrivateLocalAI")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # Separator
    doc.add_paragraph("_" * 60)

    # Document type label
    format_labels = {
        "zusammenfassung": "Zusammenfassung",
        "besprechungsprotokoll": "Besprechungsprotokoll",
        "schriftsatz_entwurf": "Schriftsatz-Entwurf",
        "mandantenanschreiben": "Mandantenanschreiben",
    }
    label = format_labels.get(document_type, "Dokument")
    title_para = doc.add_paragraph()
    run = title_para.add_run(label)
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()

    # Content: split by lines and add as paragraphs
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.font.size = Pt(14)
            run.font.bold = True
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.font.size = Pt(12)
            run.font.bold = True
        elif stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.font.size = Pt(16)
            run.font.bold = True
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.font.bold = True
        elif stripped:
            doc.add_paragraph(stripped)

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        "ENTWURF -- Erstellt durch KI-Assistenz. Pruefung erforderlich.\n"
        "PrivateLocalAI -- 100% lokale Verarbeitung"
    )
    run.font.size = Pt(8)
    run.font.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
